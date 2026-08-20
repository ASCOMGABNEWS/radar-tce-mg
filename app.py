import streamlit as st
import feedparser
from datetime import datetime, timedelta
import re
from collections import Counter
from io import BytesIO
from html import escape

from docx.shared import Pt
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.enums import TA_CENTER
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.units import cm



# ============================================================
# EXPORTAÇÃO DO CLIPPING
# ============================================================

def gerar_excel(noticias):
    linhas = []

    for n in noticias:
        linhas.append({
            "Data": n["data"].strftime("%d/%m/%Y %H:%M") if n["data"] else "",
            "Veículo": n["veiculo"],
            "Título": n["titulo"],
            "Relevância": n["bolinha"],
            "Pessoas": ", ".join(n["pessoas"]),
            "Temas": ", ".join(n["temas"]),
            "Link": n["link"],
        })

    df = pd.DataFrame(linhas)
    output = BytesIO()

    with pd.ExcelWriter(output, engine="openpyxl") as writer:
        df.to_excel(writer, index=False, sheet_name="Clipping")

    output.seek(0)
    return output


def gerar_word(noticias):
    doc = Document()

    titulo = doc.add_heading(
        f"CLIPPING TCE-MG — {datetime.now().strftime('%d/%m/%Y')}",
        0
    )
    doc.add_paragraph(
        f"{len(noticias)} notícia(s) selecionada(s)"
    )

    doc.add_heading("Destaques", level=1)

    for n in noticias:
        p = doc.add_paragraph()

        run = p.add_run(
            f"{n['bolinha']} {n['titulo']}"
        )
        run.bold = True
        run.font.size = Pt(12)

        doc.add_paragraph(
            f"{n['veiculo']} • "
            f"{n['data'].strftime('%d/%m/%Y %H:%M') if n['data'] else ''}"
        )

        if n["pessoas"]:
            doc.add_paragraph(
                "Pessoas: " + ", ".join(n["pessoas"])
            )

        if n["temas"]:
            doc.add_paragraph(
                "Temas: " + ", ".join(n["temas"])
            )

        if n["resumo"]:
            doc.add_paragraph(n["resumo"])

        doc.add_paragraph(
            "Link: " + n["link"]
        )

    output = BytesIO()
    doc.save(output)
    output.seek(0)
    return output


def gerar_pdf(noticias):
    output = BytesIO()

    doc = SimpleDocTemplate(
        output,
        pagesize=A4,
        rightMargin=1.7 * cm,
        leftMargin=1.7 * cm,
        topMargin=1.5 * cm,
        bottomMargin=1.5 * cm,
    )

    styles = getSampleStyleSheet()

    titulo_style = ParagraphStyle(
        "TituloClipping",
        parent=styles["Title"],
        alignment=TA_CENTER,
        fontSize=18,
        leading=22,
        spaceAfter=10,
    )

    noticia_style = ParagraphStyle(
        "Noticia",
        parent=styles["BodyText"],
        fontSize=10,
        leading=14,
        spaceAfter=5,
    )

    meta_style = ParagraphStyle(
        "Meta",
        parent=styles["BodyText"],
        fontSize=8,
        leading=11,
        spaceAfter=7,
    )

    story = []

    story.append(
        Paragraph(
            f"CLIPPING TCE-MG — {datetime.now().strftime('%d/%m/%Y')}",
            titulo_style
        )
    )

    story.append(
        Paragraph(
            f"{len(noticias)} notícia(s) selecionada(s)",
            meta_style
        )
    )

    for n in noticias:

        data = (
            n["data"].strftime("%d/%m/%Y %H:%M")
            if n["data"] else ""
        )

        titulo = escape(
            f"{n['bolinha']} {n['titulo']}"
        )

        story.append(
            Paragraph(
                f"<b>{titulo}</b>",
                noticia_style
            )
        )

        story.append(
            Paragraph(
                escape(
                    f"{n['veiculo']} • {data}"
                ),
                meta_style
            )
        )

        if n["pessoas"]:
            story.append(
                Paragraph(
                    escape(
                        "Pessoas: "
                        + ", ".join(n["pessoas"])
                    ),
                    meta_style
                )
            )

        if n["temas"]:
            story.append(
                Paragraph(
                    escape(
                        "Temas: "
                        + ", ".join(n["temas"])
                    ),
                    meta_style
                )
            )

        if n["resumo"]:
            resumo = n["resumo"]
            if len(resumo) > 700:
                resumo = resumo[:700] + "..."

            story.append(
                Paragraph(
                    escape(resumo),
                    noticia_style
                )
            )

        story.append(
            Paragraph(
                escape(n["link"]),
                meta_style
            )
        )

        story.append(Spacer(1, 0.22 * cm))

    doc.build(story)

    output.seek(0)
    return output


# ============================================================
# CONFIGURAÇÃO
# ============================================================

st.set_page_config(
    page_title="Radar TCE-MG",
    page_icon="🏛️",
    layout="wide"
)


# ============================================================
# FONTES
# ============================================================

FONTES = {

    # --------------------------------------------------------
    # IMPRENSA MINEIRA
    # --------------------------------------------------------

    "Estado de Minas":
        'https://news.google.com/rss/search?q=site%3Aem.com.br+%22TCE-MG%22&hl=pt-BR&gl=BR&ceid=BR:pt-419',

    "Itatiaia":
        'https://news.google.com/rss/search?q=site%3Aitatiaia.com.br+%22TCE-MG%22&hl=pt-BR&gl=BR&ceid=BR:pt-419',

    "O TEMPO":
        'https://news.google.com/rss/search?q=site%3Aotempo.com.br+%22TCE-MG%22&hl=pt-BR&gl=BR&ceid=BR:pt-419',

    "Hoje em Dia":
        'https://news.google.com/rss/search?q=site%3Ahojeemdia.com.br+%22TCE-MG%22&hl=pt-BR&gl=BR&ceid=BR:pt-419',

    "Tribuna de Minas":
        'https://news.google.com/rss/search?q=site%3Atribunademinas.com.br+%22TCE-MG%22&hl=pt-BR&gl=BR&ceid=BR:pt-419',

    "Diário do Comércio":
        'https://news.google.com/rss/search?q=site%3Adiariodocomercio.com.br+%22TCE-MG%22&hl=pt-BR&gl=BR&ceid=BR:pt-419',

    "BHAZ":
        'https://news.google.com/rss/search?q=site%3Abhaz.com.br+%22TCE-MG%22&hl=pt-BR&gl=BR&ceid=BR:pt-419',

    "Agência Minas":
        'https://news.google.com/rss/search?q=site%3Aagenciaminas.mg.gov.br+%22Tribunal%20de%20Contas%22&hl=pt-BR&gl=BR&ceid=BR:pt-419',

    "O Fator":
        'https://news.google.com/rss/search?q=site%3Aofator.com.br+%22TCE-MG%22&hl=pt-BR&gl=BR&ceid=BR:pt-419',

    "Edição do Brasil":
        'https://news.google.com/rss/search?q=site%3Aedicaodobrasil.com.br+%22TCE-MG%22&hl=pt-BR&gl=BR&ceid=BR:pt-419',

    "Moon BH":
        'https://news.google.com/rss/search?q=site%3Amoonbh.com.br+%22TCE-MG%22&hl=pt-BR&gl=BR&ceid=BR:pt-419',


    # --------------------------------------------------------
    # IMPRENSA NACIONAL
    # --------------------------------------------------------

    "G1 Minas":
        'https://news.google.com/rss/search?q=site%3Ag1.globo.com%2Fmg+%22TCE-MG%22&hl=pt-BR&gl=BR&ceid=BR:pt-419',

    "Folha de S.Paulo":
        'https://news.google.com/rss/search?q=site%3Afolha.uol.com.br+%22TCE-MG%22&hl=pt-BR&gl=BR&ceid=BR:pt-419',

    "O Globo":
        'https://news.google.com/rss/search?q=site%3Aoglobo.globo.com+%22TCE-MG%22&hl=pt-BR&gl=BR&ceid=BR:pt-419',

    "Correio Braziliense":
        'https://news.google.com/rss/search?q=site%3Acorreiobraziliense.com.br+%22TCE-MG%22&hl=pt-BR&gl=BR&ceid=BR:pt-419',

    "Poder360":
        'https://news.google.com/rss/search?q=site%3Apoder360.com.br+%22TCE-MG%22&hl=pt-BR&gl=BR&ceid=BR:pt-419',

    "O Antagonista":
        'https://news.google.com/rss/search?q=site%3Aoantagonista.com.br+%22TCE-MG%22&hl=pt-BR&gl=BR&ceid=BR:pt-419',

    "Brasil de Fato":
        'https://news.google.com/rss/search?q=site%3Abrasildefato.com.br+%22TCE-MG%22&hl=pt-BR&gl=BR&ceid=BR:pt-419',

    "revista piauí":
        'https://news.google.com/rss/search?q=site%3Apiaui.folha.uol.com.br+%22TCE-MG%22&hl=pt-BR&gl=BR&ceid=BR:pt-419',

    "JOTA":
        'https://news.google.com/rss/search?q=site%3Ajota.info+%22TCE-MG%22&hl=pt-BR&gl=BR&ceid=BR:pt-419',

    "Migalhas":
        'https://news.google.com/rss/search?q=site%3Amigalhas.com.br+%22TCE-MG%22&hl=pt-BR&gl=BR&ceid=BR:pt-419',

    "O Bastidor":
        'https://news.google.com/rss/search?q=site%3Aobastidor.com.br+%22TCE-MG%22&hl=pt-BR&gl=BR&ceid=BR:pt-419',

    "Intercept Brasil":
        'https://news.google.com/rss/search?q=site%3Aintercept.com.br+%22TCE-MG%22&hl=pt-BR&gl=BR&ceid=BR:pt-419',

    "Bem Minas":
        'https://news.google.com/rss/search?q=site%3Abemminas.com.br+%22TCE-MG%22&hl=pt-BR&gl=BR&ceid=BR:pt-419',


    # --------------------------------------------------------
    # BUSCAS GERAIS
    # --------------------------------------------------------

    "TCE-MG":
        'https://news.google.com/rss/search?q=%22TCE-MG%22&hl=pt-BR&gl=BR&ceid=BR:pt-419',

    "TCE MG":
        'https://news.google.com/rss/search?q=%22TCE%20MG%22&hl=pt-BR&gl=BR&ceid=BR:pt-419',

    "Tribunal de Contas MG":
        'https://news.google.com/rss/search?q=%22Tribunal%20de%20Contas%22+%22Minas%20Gerais%22&hl=pt-BR&gl=BR&ceid=BR:pt-419',


    # --------------------------------------------------------
    # PESSOAS
    # --------------------------------------------------------

    "Durval Ângelo":
        'https://news.google.com/rss/search?q=%22Durval%20Ângelo%22&hl=pt-BR&gl=BR&ceid=BR:pt-419',

    "Agostinho Patrus":
        'https://news.google.com/rss/search?q=%22Agostinho%20Patrus%22&hl=pt-BR&gl=BR&ceid=BR:pt-419',

    "Gilberto Diniz":
        'https://news.google.com/rss/search?q=%22Gilberto%20Diniz%22+TCE&hl=pt-BR&gl=BR&ceid=BR:pt-419',

    "Alencar da Silveira":
        'https://news.google.com/rss/search?q=%22Alencar%20da%20Silveira%22+TCE&hl=pt-BR&gl=BR&ceid=BR:pt-419',

    "Ione Pinheiro":
        'https://news.google.com/rss/search?q=%22Ione%20Pinheiro%22+TCE&hl=pt-BR&gl=BR&ceid=BR:pt-419',

    "Tadeu Martins Leite / Tadeuzinho":
        'https://news.google.com/rss/search?q=%22Tadeu%20Martins%20Leite%22+OR+%22Tadeuzinho%22&hl=pt-BR&gl=BR&ceid=BR:pt-419',

    "Licurgo Mourão":
        'https://news.google.com/rss/search?q=%22Licurgo%20Mourão%22+TCE&hl=pt-BR&gl=BR&ceid=BR:pt-419',

    "Hamilton Coelho":
        'https://news.google.com/rss/search?q=%22Hamilton%20Coelho%22+TCE&hl=pt-BR&gl=BR&ceid=BR:pt-419',

    "Adonias Fernandes":
        'https://news.google.com/rss/search?q=%22Adonias%20Fernandes%22+TCE&hl=pt-BR&gl=BR&ceid=BR:pt-419',

    "Telmo Passareli":
        'https://news.google.com/rss/search?q=%22Telmo%20Passareli%22+TCE&hl=pt-BR&gl=BR&ceid=BR:pt-419',
}


# ============================================================
# PESSOAS MONITORADAS
# ============================================================

PESSOAS = {

    "Conselheiros": {

        "Durval Ângelo": [
            "Durval Ângelo"
        ],

        "Agostinho Patrus": [
            "Agostinho Patrus"
        ],

        "Gilberto Diniz": [
            "Gilberto Diniz"
        ],

        "Alencar da Silveira": [
            "Alencar da Silveira",
            "Alencar Silveira"
        ],

        "Ione Pinheiro": [
            "Ione Pinheiro"
        ],
    },

    "Eleito / transição": {

        "Tadeu Martins Leite (Tadeuzinho)": [
            "Tadeu Martins Leite",
            "Tadeuzinho",
            "Tadeu Leite"
        ],
    },

    "Conselheiros substitutos": {

        "Licurgo Joseph Mourão de Oliveira": [
            "Licurgo Joseph Mourão",
            "Licurgo Mourão"
        ],

        "Hamilton Antônio Coelho": [
            "Hamilton Antônio Coelho",
            "Hamilton Coelho"
        ],

        "Adonias Fernandes Monteiro": [
            "Adonias Fernandes Monteiro",
            "Adonias Fernandes"
        ],

        "Telmo de Moura Passareli": [
            "Telmo de Moura Passareli",
            "Telmo Passareli"
        ],
    }
}


# ============================================================
# RELAÇÃO BUSCA → PESSOA
# ============================================================

MAPA_FONTE_PESSOA = {

    "Durval Ângelo":
        "Durval Ângelo",

    "Agostinho Patrus":
        "Agostinho Patrus",

    "Gilberto Diniz":
        "Gilberto Diniz",

    "Alencar da Silveira":
        "Alencar da Silveira",

    "Ione Pinheiro":
        "Ione Pinheiro",

    "Tadeu Martins Leite / Tadeuzinho":
        "Tadeu Martins Leite (Tadeuzinho)",

    "Licurgo Mourão":
        "Licurgo Joseph Mourão de Oliveira",

    "Hamilton Coelho":
        "Hamilton Antônio Coelho",

    "Adonias Fernandes":
        "Adonias Fernandes Monteiro",

    "Telmo Passareli":
        "Telmo de Moura Passareli",
}


# ============================================================
# TEMAS
# ============================================================

TEMAS = {

    "⛏️ Mineração": [
        "mineração",
        "mineradora",
        "Vale",
        "CSN",
        "CFEM",
        "royalties",
        "barragem",
        "barragens"
    ],

    "🚰 Copasa": [
        "Copasa",
        "saneamento",
        "água",
        "abastecimento"
    ],

    "🏭 Estatais": [
        "Cemig",
        "Codemig",
        "empresa estatal",
        "estatal"
    ],

    "🏛️ Privatização": [
        "privatização",
        "privatizar",
        "desestatização",
        "venda da estatal"
    ],

    "🚌 Transporte": [
        "transporte",
        "ônibus",
        "metro",
        "metrô",
        "rodovia",
        "pedágio",
        "concessão",
        "mobilidade"
    ],

    "💰 Benefícios fiscais": [
        "benefícios fiscais",
        "benefício fiscal",
        "incentivos fiscais",
        "incentivo fiscal",
        "renúncia fiscal",
        "ICMS",
        "crédito presumido"
    ],

    "⚖️ Controle": [
        "auditoria",
        "fiscalização",
        "licitação",
        "contrato",
        "contas",
        "denúncia",
        "irregularidade",
        "multa",
        "ressarcimento",
        "julgamento",
        "acórdão",
        "determina"
    ],

    "🏥 Saúde": [
        "saúde",
        "hospital",
        "hospitais",
        "SUS",
        "medicamento"
    ],

    "🎓 Educação": [
        "educação",
        "escola",
        "escolas",
        "ensino",
        "universidade",
        "educacional"
    ],

    "🏗️ Obras públicas": [
        "obra pública",
        "obras públicas",
        "obras",
        "infraestrutura",
        "construção"
    ],
}


# ============================================================
# FUNÇÕES
# ============================================================

def limpar_texto(texto):

    if not texto:
        return ""

    texto = re.sub(
        r"<[^>]+>",
        " ",
        texto
    )

    texto = re.sub(
        r"\s+",
        " ",
        texto
    )

    return texto.strip()


def obter_data(item):

    try:

        if (
            hasattr(item, "published_parsed")
            and item.published_parsed
        ):

            return datetime(
                *item.published_parsed[:6]
            )

    except Exception:
        pass

    return None


def extrair_veiculo(item):

    titulo = item.get(
        "title",
        ""
    )

    if " - " in titulo:

        return titulo.rsplit(
            " - ",
            1
        )[-1].strip()

    return "Fonte não identificada"


def identificar_temas(
    titulo,
    resumo
):

    texto = (
        titulo
        + " "
        + resumo
    ).lower()

    encontrados = []

    for tema, palavras in TEMAS.items():

        for palavra in palavras:

            if palavra.lower() in texto:

                encontrados.append(
                    tema
                )

                break

    return encontrados


def identificar_pessoas(
    titulo,
    resumo
):

    texto = (
        titulo
        + " "
        + resumo
    ).lower()

    encontrados = []

    for grupo in PESSOAS.values():

        for pessoa, variacoes in grupo.items():

            for variacao in variacoes:

                if (
                    variacao.lower()
                    in texto
                ):

                    if pessoa not in encontrados:

                        encontrados.append(
                            pessoa
                        )

                    break

    return encontrados


def calcular_relevancia(
    titulo,
    resumo,
    monitoramento,
    temas,
    pessoas
):

    texto = (
        titulo
        + " "
        + resumo
    ).lower()

    score = 15


    if "tce-mg" in texto:

        score += 35

    elif "tce mg" in texto:

        score += 30

    elif "tribunal de contas" in texto:

        score += 25


    score += (
        len(pessoas) * 12
    )


    score += (
        len(temas) * 5
    )


    termos_acao = [

        "determina",
        "decide",
        "suspende",
        "condena",
        "multa",
        "auditoria",
        "fiscalização",
        "julgamento",
        "acórdão",
        "denúncia",
        "irregularidade",
        "recomenda",
        "processo",
        "ressarcimento",
        "contas",
    ]


    for termo in termos_acao:

        if termo in texto:

            score += 5


    if "r$" in texto:

        score += 5


    return min(
        score,
        100
    )


def classificar(score):

    if score >= 85:

        return "🔴"

    if score >= 65:

        return "🟠"

    if score >= 45:

        return "🟡"

    return "⚪"


# ============================================================
# COLETA
# ============================================================

@st.cache_data(ttl=300)
def buscar_noticias():

    noticias = []

    links = set()

    limite = (
        datetime.now()
        - timedelta(days=7)
    )


    for nome, url in FONTES.items():

        try:

            feed = feedparser.parse(
                url
            )

        except Exception:

            continue


        for item in feed.entries:

            link = item.get(
                "link",
                ""
            )


            if (
                not link
                or link in links
            ):

                continue


            data = obter_data(
                item
            )


            if (
                data
                and data < limite
            ):

                continue


            links.add(
                link
            )


            titulo = item.get(
                "title",
                "Sem título"
            )


            resumo = limpar_texto(
                item.get(
                    "summary",
                    ""
                )
            )


            temas = identificar_temas(
                titulo,
                resumo
            )


            pessoas = identificar_pessoas(
                titulo,
                resumo
            )


            # ------------------------------------------------
            # CORREÇÃO DAS PESSOAS
            # ------------------------------------------------

            if nome in MAPA_FONTE_PESSOA:

                pessoa_fonte = (
                    MAPA_FONTE_PESSOA[
                        nome
                    ]
                )

                if pessoa_fonte not in pessoas:

                    pessoas.append(
                        pessoa_fonte
                    )


            score = calcular_relevancia(

                titulo,
                resumo,
                nome,
                temas,
                pessoas
            )


            noticias.append({

                "titulo":
                    titulo,

                "resumo":
                    resumo,

                "link":
                    link,

                "monitoramento":
                    nome,

                "veiculo":
                    extrair_veiculo(
                        item
                    ),

                "data":
                    data,

                "score":
                    score,

                "bolinha":
                    classificar(
                        score
                    ),

                "temas":
                    temas,

                "pessoas":
                    pessoas,
            })


    noticias.sort(

        key=lambda x: (

            x["score"],

            x["data"]
            or datetime.min

        ),

        reverse=True
    )


    return noticias


# ============================================================
# CABEÇALHO
# ============================================================

st.title(
    "🏛️ Radar TCE-MG"
)

st.caption(
    "Monitoramento inteligente de notícias "
    "relacionadas ao Tribunal de Contas de Minas Gerais"
)


col1, col2 = st.columns(
    [1, 5]
)


with col1:

    if st.button(
        "🔄 Atualizar agora"
    ):

        st.cache_data.clear()

        st.rerun()


with col2:

    st.caption(
        "Monitoramento em tempo quase real • "
        "atualização automática a cada 5 minutos"
    )


st.divider()


# ============================================================
# COLETA
# ============================================================

noticias = buscar_noticias()


# ============================================================
# PERÍODO
# ============================================================

periodo = st.radio(

    "🕐 Período",

    [
        "Últimas 6 horas",
        "Últimas 24 horas",
        "Últimos 3 dias",
        "Últimos 7 dias",
    ],

    horizontal=True
)


agora = datetime.now()


if periodo == "Últimas 6 horas":

    limite_periodo = (
        agora
        - timedelta(hours=6)
    )

elif periodo == "Últimas 24 horas":

    limite_periodo = (
        agora
        - timedelta(hours=24)
    )

elif periodo == "Últimos 3 dias":

    limite_periodo = (
        agora
        - timedelta(days=3)
    )

else:

    limite_periodo = (
        agora
        - timedelta(days=7)
    )


noticias_periodo = [

    n for n in noticias

    if (
        n["data"]
        and n["data"]
        >= limite_periodo
    )
]


# ============================================================
# RADAR DE ATENÇÃO
# ============================================================

criticas = [

    n for n in noticias_periodo

    if n["score"] >= 85
]


altas = [

    n for n in noticias_periodo

    if 65 <= n["score"] < 85
]


st.subheader(
    "🚨 Radar de atenção"
)


col1, col2 = st.columns(
    2
)


with col1:

    st.metric(
        "🔴 Críticas",
        len(criticas)
    )


with col2:

    st.metric(
        "🟠 Alta relevância",
        len(altas)
    )


if criticas:

    st.markdown(
        "#### 🔴 Atenção imediata"
    )


    for noticia in criticas[:5]:

        st.markdown(
            f"**{noticia['titulo']}**"
        )


st.divider()


# ============================================================
# CONTADORES
# ============================================================

contador_temas = Counter()

contador_pessoas = Counter()


for noticia in noticias_periodo:

    for tema in noticia["temas"]:

        contador_temas[
            tema
        ] += 1


    for pessoa in noticia["pessoas"]:

        contador_pessoas[
            pessoa
        ] += 1


# ============================================================
# ASSUNTOS + PESSOAS
# ============================================================

col1, col2 = st.columns(
    2
)


with col1:

    st.subheader(
        "🔥 Assuntos quentes"
    )


    if contador_temas:

        temas_quentes = (
            contador_temas
            .most_common(8)
        )


        for tema, quantidade in temas_quentes:

            st.markdown(
                f"**{tema}** — "
                f"{quantidade} notícia(s)"
            )

    else:

        st.info(
            "Nenhum tema identificado."
        )


with col2:

    st.subheader(
        "👥 Pessoas mais citadas"
    )


    if contador_pessoas:

        pessoas_quentes = (
            contador_pessoas
            .most_common(8)
        )


        for pessoa, quantidade in pessoas_quentes:

            st.markdown(
                f"**{pessoa}** — "
                f"{quantidade} notícia(s)"
            )

    else:

        st.info(
            "Nenhuma pessoa monitorada citada."
        )


st.divider()


# ============================================================
# MÉTRICAS
# ============================================================

criticas_total = len([

    n for n in noticias_periodo

    if n["score"] >= 85
])


altas_total = len([

    n for n in noticias_periodo

    if 65 <= n["score"] < 85
])


medias_total = len([

    n for n in noticias_periodo

    if 45 <= n["score"] < 65
])


mencoes_total = len([

    n for n in noticias_periodo

    if n["score"] < 45
])


col1, col2, col3, col4, col5 = st.columns(
    5
)


with col1:

    st.metric(
        "📰 Notícias",
        len(noticias_periodo)
    )


with col2:

    st.metric(
        "🔴 Críticas",
        criticas_total
    )


with col3:

    st.metric(
        "🟠 Altas",
        altas_total
    )


with col4:

    st.metric(
        "🟡 Médias",
        medias_total
    )


with col5:

    st.metric(
        "⚪ Menções",
        mencoes_total
    )


st.divider()


# ============================================================
# FILTROS
# ============================================================

st.subheader(
    "🔎 Monitorar"
)


todas_pessoas = []


for grupo in PESSOAS.values():

    todas_pessoas.extend(
        grupo.keys()
    )


col1, col2, col3 = st.columns(
    3
)


with col1:

    filtro_pessoa = st.selectbox(

        "👤 Pessoa",

        [
            "Todas"
        ]
        + todas_pessoas
    )


with col2:

    filtro_tema = st.selectbox(

        "🏷️ Tema",

        [
            "Todos"
        ]
        + list(
            TEMAS.keys()
        )
    )


with col3:

    filtro_fonte = st.selectbox(

        "🗞️ Fonte",

        [
            "Todas"
        ]
        + list(
            FONTES.keys()
        )
    )


col1, col2 = st.columns(
    2
)


with col1:

    filtro_relevancia = st.selectbox(

        "🎯 Relevância",

        [
            "Todas",
            "🔴 Crítica",
            "🟠 Alta",
            "🟡 Média",
            "⚪ Menção"
        ]
    )


with col2:

    busca = st.text_input(

        "🔍 Buscar palavra",

        placeholder=
        "Ex.: Copasa, mineração, transporte..."
    )


# ============================================================
# APLICA FILTROS
# ============================================================

filtradas = noticias_periodo


if filtro_pessoa != "Todas":

    filtradas = [

        n for n in filtradas

        if filtro_pessoa
        in n["pessoas"]

    ]


if filtro_tema != "Todos":

    filtradas = [

        n for n in filtradas

        if filtro_tema
        in n["temas"]

    ]


if filtro_fonte != "Todas":

    filtradas = [

        n for n in filtradas

        if n["monitoramento"]
        == filtro_fonte

    ]


if filtro_relevancia != "Todas":

    mapa_relevancia = {

        "🔴 Crítica": "🔴",
        "🟠 Alta": "🟠",
        "🟡 Média": "🟡",
        "⚪ Menção": "⚪",
    }


    filtradas = [

        n for n in filtradas

        if n["bolinha"]
        == mapa_relevancia[
            filtro_relevancia
        ]

    ]


if busca:

    termo = busca.lower()

    filtradas = [

        n for n in filtradas

        if termo
        in (
            n["titulo"]
            + " "
            + n["resumo"]
        ).lower()

    ]



# ============================================================
# DOWNLOAD DO CLIPPING
# ============================================================

from io import BytesIO
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.enums import TA_CENTER
from reportlab.lib import colors
from reportlab.platypus import (
    SimpleDocTemplate,
    Paragraph,
    Spacer,
    PageBreak,
)
from reportlab.lib.units import cm


def gerar_pdf_clipping(noticias_clipping):
    buffer = BytesIO()

    doc = SimpleDocTemplate(
        buffer,
        pagesize=A4,
        rightMargin=1.6 * cm,
        leftMargin=1.6 * cm,
        topMargin=1.6 * cm,
        bottomMargin=1.6 * cm,
    )

    styles = getSampleStyleSheet()

    titulo = ParagraphStyle(
        "TituloClipping",
        parent=styles["Title"],
        fontSize=20,
        leading=24,
        alignment=TA_CENTER,
        spaceAfter=8,
    )

    subtitulo = ParagraphStyle(
        "SubtituloClipping",
        parent=styles["Normal"],
        fontSize=10,
        leading=14,
        alignment=TA_CENTER,
        textColor=colors.grey,
        spaceAfter=18,
    )

    secao = ParagraphStyle(
        "SecaoClipping",
        parent=styles["Heading2"],
        fontSize=13,
        leading=16,
        spaceBefore=12,
        spaceAfter=8,
    )

    manchete = ParagraphStyle(
        "MancheteClipping",
        parent=styles["Heading3"],
        fontSize=11.5,
        leading=15,
        spaceAfter=5,
    )

    corpo = ParagraphStyle(
        "CorpoClipping",
        parent=styles["Normal"],
        fontSize=9,
        leading=13,
        spaceAfter=6,
    )

    meta = ParagraphStyle(
        "MetaClipping",
        parent=styles["Normal"],
        fontSize=8,
        leading=11,
        textColor=colors.grey,
        spaceAfter=5,
    )

    story = []

    data_geracao = datetime.now().strftime("%d/%m/%Y às %H:%M")

    story.append(
        Paragraph(
            "CLIPPING TCE-MG",
            titulo,
        )
    )

    story.append(
        Paragraph(
            f"Gerado em {data_geracao} • "
            f"{len(noticias_clipping)} notícia(s)",
            subtitulo,
        )
    )

    # Destaques
    criticas_pdf = [
        n for n in noticias_clipping
        if n["score"] >= 85
    ]

    altas_pdf = [
        n for n in noticias_clipping
        if 65 <= n["score"] < 85
    ]

    if criticas_pdf or altas_pdf:

        story.append(
            Paragraph(
                "DESTAQUES",
                secao,
            )
        )

        for noticia in (
            criticas_pdf[:5] + altas_pdf[:5]
        ):

            story.append(
                Paragraph(
                    f"{noticia['bolinha']} "
                    f"{noticia['titulo']}",
                    manchete,
                )
            )

            story.append(
                Paragraph(
                    f"<b>{noticia['veiculo']}</b>",
                    meta,
                )
            )

    # Notícias
    story.append(
        Paragraph(
            "NOTÍCIAS",
            secao,
        )
    )

    for noticia in noticias_clipping:

        data = ""

        if noticia["data"]:
            data = noticia["data"].strftime(
                "%d/%m/%Y %H:%M"
            )

        pessoas = ", ".join(
            noticia["pessoas"]
        )

        temas = ", ".join(
            noticia["temas"]
        )

        resumo = noticia["resumo"]

        if len(resumo) > 700:
            resumo = resumo[:700] + "..."

        story.append(
            Paragraph(
                f"{noticia['bolinha']} "
                f"{noticia['titulo']}",
                manchete,
            )
        )

        story.append(
            Paragraph(
                f"{noticia['veiculo']} • {data}",
                meta,
            )
        )

        if pessoas:
            story.append(
                Paragraph(
                    f"<b>Pessoas:</b> {pessoas}",
                    corpo,
                )
            )

        if temas:
            story.append(
                Paragraph(
                    f"<b>Temas:</b> {temas}",
                    corpo,
                )
            )

        if resumo:
            story.append(
                Paragraph(
                    resumo,
                    corpo,
                )
            )

        story.append(
            Paragraph(
                f'<link href="{noticia["link"]}" '
                f'color="blue">{noticia["link"]}</link>',
                meta,
            )
        )

        story.append(
            Spacer(
                1,
                0.18 * cm
            )
        )

    # Assuntos
    contador_temas_pdf = Counter()

    for noticia in noticias_clipping:
        for tema in noticia["temas"]:
            contador_temas_pdf[tema] += 1

    if contador_temas_pdf:

        story.append(
            PageBreak()
        )

        story.append(
            Paragraph(
                "ASSUNTOS EM DESTAQUE",
                secao,
            )
        )

        for tema, quantidade in (
            contador_temas_pdf.most_common(10)
        ):

            story.append(
                Paragraph(
                    f"{tema} — {quantidade} notícia(s)",
                    corpo,
                )
            )

    # Pessoas
    contador_pessoas_pdf = Counter()

    for noticia in noticias_clipping:
        for pessoa in noticia["pessoas"]:
            contador_pessoas_pdf[pessoa] += 1

    if contador_pessoas_pdf:

        story.append(
            Paragraph(
                "PESSOAS MAIS CITADAS",
                secao,
            )
        )

        for pessoa, quantidade in (
            contador_pessoas_pdf.most_common(10)
        ):

            story.append(
                Paragraph(
                    f"{pessoa} — {quantidade} notícia(s)",
                    corpo,
                )
            )

    doc.build(story)

    buffer.seek(0)

    return buffer.getvalue()


pdf_bytes = gerar_pdf_clipping(
    filtradas
)

st.download_button(
    label="📄 Baixar clipping em PDF",
    data=pdf_bytes,
    file_name=(
        f"clipping_tce_mg_"
        f"{datetime.now().strftime('%Y-%m-%d')}.pdf"
    ),
    mime="application/pdf",
)
